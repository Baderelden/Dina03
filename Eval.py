import io
from datetime import datetime

import streamlit as st

from langchain_openai import ChatOpenAI
from langchain.prompts import ChatPromptTemplate

from pypdf import PdfReader
import docx
from docx import Document


# ---------- Helpers to read files ----------

def read_uploaded_file(uploaded_file) -> str:
    """
    Returns the textual content of the uploaded file as a string.
    Supports: .txt, .md, .pdf, .docx (basic support).
    """
    filename = uploaded_file.name.lower()

    if filename.endswith((".txt", ".md")):
        return uploaded_file.read().decode("utf-8", errors="ignore")

    if filename.endswith(".pdf"):
        reader = PdfReader(uploaded_file)
        text = []
        for page in reader.pages:
            text.append(page.extract_text() or "")
        return "\n".join(text)

    if filename.endswith(".docx"):
        doc = docx.Document(uploaded_file)
        return "\n".join(p.text for p in doc.paragraphs)

    # Fallback: try to decode as text
    try:
        return uploaded_file.read().decode("utf-8", errors="ignore")
    except Exception:
        return ""


# ---------- Prompt helpers ----------

def build_aspect_instructions(aspect_choice: str) -> str:
    if aspect_choice == "Alignment with Kuwait 2035 Vision":
        return (
            "Evaluate how well this research proposal aligns with Kuwait Vision 2035 "
            "priorities (e.g., knowledge economy, innovation, human capital development, "
            "sustainability, digital transformation). Identify strengths, weaknesses, "
            "and any missing links."
        )
    elif aspect_choice == "Alignment with AASU vision & mission":
        return (
            "Evaluate how well this proposal aligns with a modern research-led university vision, "
            "such as AASU’s focus on high-quality education, impactful research, innovation, "
            "and service to society. Comment on fit with likely strategic themes, reputation building, "
            "and benefits to students and staff."
        )
    elif aspect_choice == "Ethical considerations & implications":
        return (
            "Evaluate the ethical aspects of this proposal. Consider research ethics, data protection, "
            "participant welfare, inclusivity, fairness, and any potential negative societal impacts. "
            "Highlight areas that require ethics committee attention or clearer mitigation plans."
        )
    elif aspect_choice == "Capacity building at AASU":
        return (
            "Evaluate the extent to which this proposal contributes to capacity building at AASU. "
            "Consider staff development, student training, new research capabilities, laboratories, "
            "international collaborations, and long-term benefits for institutional growth."
        )
    elif aspect_choice == "All of the above":
        return (
            "Evaluate this proposal from FOUR angles:\n"
            "1) Alignment with Kuwait Vision 2035 priorities.\n"
            "2) Alignment with AASU’s overall vision, mission, and strategic direction.\n"
            "3) Ethical considerations and potential risks or implications.\n"
            "4) Contribution to capacity building at AASU (people, infrastructure, partnerships, reputation).\n"
            "Provide a structured response with clear sections."
        )
    else:
        return "Provide a general evaluation of this research proposal."


def build_detail_instruction(detail_level: str) -> str:
    if detail_level == "Short summary":
        return (
            "Keep your answer concise, around 2–4 short paragraphs, focusing on the most important points only."
        )
    else:  # "Detailed report"
        return (
            "Provide a detailed, structured evaluation, with headings, bullet points where useful, "
            "and specific suggestions for improvement."
        )


# ---------- LLM factory ----------

def get_llm(model_name: str):
    return ChatOpenAI(
        model=model_name,
        temperature=0.3,
    )


# ---------- Report generation (DOCX) ----------

def create_report_docx(
    proposal_filename: str,
    model_name: str,
    aspect_choice: str,
    detail_level: str,
    extra_instructions: str,
    evaluation_text: str,
) -> io.BytesIO:
    """
    Create a .docx report in memory and return it as a BytesIO buffer.
    """
    doc = Document()

    # Title
    doc.add_heading("AASU Research Proposal Evaluation", level=1)

    # Meta information
    doc.add_paragraph(f"Date: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    doc.add_paragraph(f"Proposal file: {proposal_filename}")
    doc.add_paragraph(f"Model (engine): {model_name}")
    doc.add_paragraph(f"Evaluation focus: {aspect_choice}")
    doc.add_paragraph(f"Detail level: {detail_level}")

    if extra_instructions.strip():
        doc.add_paragraph("Extra instructions from user:")
        doc.add_paragraph(extra_instructions, style="Intense Quote")

    doc.add_paragraph("")  # spacing

    # Evaluation body
    doc.add_heading("Evaluation", level=2)
    for line in evaluation_text.split("\n"):
        # Preserve some structure but avoid creating a heading for every line
        if line.strip().startswith("#"):
            # If user / model used markdown headings, make them subheadings
            clean = line.lstrip("#").strip()
            if clean:
                doc.add_heading(clean, level=3)
        else:
            doc.add_paragraph(line)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# ---------- Streamlit UI ----------

def main():
    st.set_page_config(
        page_title="AASU Research Proposal Evaluator",
        layout="wide",
    )

    st.title("AASU Research Proposal Evaluator")
    st.write(
        "Upload a research proposal and generate an evaluation based on AASU and Kuwait 2035 priorities."
    )

    # --- Model / engine selection ---
    st.subheader("Model settings")

    model_name = st.selectbox(
        "Choose the language model (engine):",
        options=[
            "gpt-4.1",
            "gpt-4o",
            "gpt-4o-mini",
            "gpt-4.1-mini",
        ],
        index=2,  # default: gpt-4o-mini
        help="Pick the engine you want to use for the evaluation.",
    )

    st.caption(
        "- gpt-4o / gpt-4.1: higher quality, more expensive, a bit slower.\n"
        "- gpt-4o-mini / gpt-4.1-mini: cheaper, faster, good for quick checks."
    )

    st.markdown("---")

    # 1) File upload
    uploaded_file = st.file_uploader(
        "Upload research proposal file",
        type=["pdf", "docx", "txt", "md"],
        help="Supported formats: PDF, DOCX, TXT, MD",
    )

    st.markdown("---")

    # 2) Aspect selection
    aspect_choice = st.radio(
        "What would you like the app to check?",
        options=[
            "Alignment with Kuwait 2035 Vision",
            "Alignment with AASU vision & mission",
            "Ethical considerations & implications",
            "Capacity building at AASU",
            "All of the above",
        ],
        index=4,  # default: All of the above
    )

    # 3) Detail level
    detail_level = st.radio(
        "Level of detail in the answer:",
        options=["Short summary", "Detailed report"],
        index=0,
    )

    # 4) Extra instructions
    extra_instructions = st.text_area(
        "Extra instructions (optional)",
        placeholder=(
            "e.g. ‘Focus on AI ethics’, ‘Assume the reviewers are non-experts’, "
            "‘Highlight risks related to data protection’, etc."
        ),
        height=120,
    )

    st.markdown("---")

    evaluation_text = None  # keep for download section

    # 5) Submit
    if st.button("Evaluate proposal", type="primary"):
        if uploaded_file is None:
            st.error("Please upload a proposal file before submitting.")
            return

        with st.spinner("Reading file and generating evaluation..."):
            proposal_text = read_uploaded_file(uploaded_file)

            if not proposal_text.strip():
                st.error(
                    "Could not extract text from this file. "
                    "Please try a different format (e.g. PDF with selectable text, DOCX, or TXT)."
                )
                return

            aspect_instructions = build_aspect_instructions(aspect_choice)
            detail_instruction = build_detail_instruction(detail_level)

            system_prompt = (
                "You are an experienced academic reviewer at Abdullah Al-Salem University (AASU). "
                "You evaluate research proposals for alignment with institutional strategies, "
                "Kuwait national priorities, ethics, and capacity building.\n\n"
                "Always be constructive, specific, and professional. When information is missing, "
                "explicitly say what is missing instead of inventing details. Give score out of 100% for each section"
            )

            template = ChatPromptTemplate.from_messages(
                [
                    ("system", system_prompt),
                    (
                        "user",
                        (
                            "You are given a research proposal. Your task:\n"
                            "{aspect_instructions}\n\n"
                            "{detail_instruction}\n\n"
                            "If extra instructions are provided, follow them as well:\n"
                            "Extra instructions from the user: {extra_instructions}\n\n"
                            "Here is the proposal text:\n"
                            "-------------------------\n"
                            "{proposal_text}\n"
                            "-------------------------\n"
                        ),
                    ),
                ]
            )

            llm = get_llm(model_name)
            chain = template | llm

            response = chain.invoke(
                {
                    "aspect_instructions": aspect_instructions,
                    "detail_instruction": detail_instruction,
                    "extra_instructions": extra_instructions or "None.",
                    "proposal_text": proposal_text,
                }
            )

            evaluation_text = response.content

        st.subheader("Evaluation")
        st.markdown(evaluation_text)

        # ---------- Download section ----------
        st.markdown("---")
        st.subheader("Download report")

        report_buffer = create_report_docx(
            proposal_filename=uploaded_file.name,
            model_name=model_name,
            aspect_choice=aspect_choice,
            detail_level=detail_level,
            extra_instructions=extra_instructions,
            evaluation_text=evaluation_text,
        )

        timestamp = datetime.now().strftime("%Y%m%d_%H%M")
        default_filename = f"AASU_Proposal_Evaluation_{timestamp}.docx"

        st.download_button(
            label="Download evaluation report (.docx)",
            data=report_buffer,
            file_name=default_filename,
            mime=(
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            ),
        )


if __name__ == "__main__":
    main()

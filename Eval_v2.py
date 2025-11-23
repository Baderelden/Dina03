import io
from datetime import datetime

import streamlit as st

from langchain_openai import ChatOpenAI
from langchain_core.prompts import ChatPromptTemplate

from pypdf import PdfReader
import docx
from docx import Document


# ---------- Helpers to read files ----------

def read_uploaded_file(uploaded_file) -> str:
    """
    Returns the textual content of the uploaded file as a string.
    Supports: .txt, .md, .pdf, .docx (basic support).
    from langchain.prompts import ChatPromptTemplate
    """
    filename = uploaded_file.name.lower()

    if filename.endswith((".txt", ".md")):
        return uploaded_file.read().decode("utf-8", errors="ignore")

    if filename.endswith(".pdf"):
        reader = PdfReader(uploaded_file)
        text = []
        for page in reader.pages:
            text.append(page.extract_text() or "")
        return "\n".join(text)

    if filename.endswith(".docx"):
        doc = docx.Document(uploaded_file)
        return "\n".join(p.text for p in doc.paragraphs)

    # Fallback: try to decode as text
    try:
        return uploaded_file.read().decode("utf-8", errors="ignore")
    except Exception:
        return ""


# ---------- Prompt helpers ----------

def build_aspect_instructions(aspect_choice: str, language: str) -> str:
    # System logic stays in English for reliability; we adjust only style / phrasing slightly if Arabic is chosen.
    if aspect_choice == "kuwait":
        return (
            "Evaluate how well this research proposal aligns with Kuwait Vision 2035 "
            "priorities such as knowledge economy, innovation, human capital development, "
            "sustainability, and digital transformation. Identify strengths, weaknesses, "
            "and any missing links."
        )
    elif aspect_choice == "aasu":
        return (
            "Evaluate how well this proposal aligns with a modern research-led university vision, "
            "such as AASU’s focus on high-quality education, impactful research, innovation, "
            "and service to society. Comment on fit with strategic themes, reputation building, "
            "and benefits to students and staff."
        )
    elif aspect_choice == "ethics":
        return (
            "Evaluate the ethical aspects of this proposal. Consider research ethics, data protection, "
            "participant welfare, inclusivity, fairness, and any potential negative societal impacts. "
            "Highlight areas that require ethics committee attention or clearer mitigation plans."
        )
    elif aspect_choice == "capacity":
        return (
            "Evaluate the extent to which this proposal contributes to capacity building at AASU. "
            "Consider staff development, student training, new research capabilities, laboratories, "
            "international collaborations, and long-term benefits for institutional growth."
        )
    elif aspect_choice == "all":
        return (
            "Evaluate this proposal from four angles:\n"
            "1) Alignment with Kuwait Vision 2035 priorities.\n"
            "2) Alignment with AASU’s overall vision, mission, and strategic direction.\n"
            "3) Ethical considerations and potential risks or implications.\n"
            "4) Contribution to capacity building at AASU (people, infrastructure, partnerships, reputation).\n"
            "Provide a structured response with clear sections."
        )
    else:
        return "Provide a general evaluation of this research proposal."


def build_detail_instruction(detail_level: str, language: str) -> str:
    if detail_level == "short":
        return (
            "Keep your answer concise, around 2–4 short paragraphs, focusing on the most important points only."
        )
    else:  # "detailed"
        return (
            "Provide a detailed, structured evaluation, with headings, bullet points where useful, "
            "and specific suggestions for improvement."
        )


def build_language_instruction(language: str) -> str:
    if language == "Arabic":
        return "Write your full evaluation in clear, formal Arabic suitable for academic committees."
    else:
        return "Write your full evaluation in clear, formal British English suitable for academic committees."


# ---------- LLM factory ----------

def get_llm(model_name: str):
    return ChatOpenAI(
        model=model_name,
        temperature=0.3,
    )


# ---------- Report generation (DOCX) ----------

def create_report_docx(
    proposal_filename: str,
    model_name: str,
    aspect_label: str,
    detail_label: str,
    extra_instructions: str,
    evaluation_text: str,
    language: str,
) -> io.BytesIO:
    """
    Create a .docx report in memory and return it as a BytesIO buffer.
    """
    doc = Document()

    # Title
    doc.add_heading("Abdullah Al-Salem University", level=1)
    doc.add_heading("Research Proposal Evaluation Report", level=2)

    # Meta information
    doc.add_paragraph(f"Date: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    doc.add_paragraph(f"Proposal file: {proposal_filename}")
    doc.add_paragraph(f"Language: {language}")
    doc.add_paragraph(f"Model (engine): {model_name}")
    doc.add_paragraph(f"Evaluation focus: {aspect_label}")
    doc.add_paragraph(f"Detail level: {detail_label}")

    if extra_instructions.strip():
        doc.add_paragraph("Extra instructions from user:")
        doc.add_paragraph(extra_instructions, style="Intense Quote")

    doc.add_paragraph("")  # spacing

    # Evaluation body
    doc.add_heading("Evaluation", level=2)
    for line in evaluation_text.split("\n"):
        # Basic handling of markdown-style headings
        if line.strip().startswith("#"):
            clean = line.lstrip("#").strip()
            if clean:
                doc.add_heading(clean, level=3)
        else:
            doc.add_paragraph(line)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# ---------- Local CSS for a formal AASU look (blue + orange) ----------

def inject_css():
    st.markdown(
        """
        <style>
        body {
            background-color: #f5f5f7;
        }

        .main-block {
            background-color: #ffffff;
            padding: 2.5rem 3rem;
            border-radius: 12px;
            border: 1px solid #dde1e7;
            box-shadow: 0 2px 8px rgba(0, 0, 0, 0.03);
            margin-bottom: 2rem;
        }

        h1, h2, h3 {
            font-family: "Georgia", "Times New Roman", serif;
            color: #002B5C;
        }

        .stMarkdown, .stRadio, .stTextArea, .stSelectbox, .stFileUploader {
            font-family: "Segoe UI", system-ui, -apple-system, BlinkMacSystemFont, sans-serif;
        }

        .stButton>button {
            background-color: #002B5C !important;
            color: white !important;
            border-radius: 6px;
            padding: 0.55rem 1.2rem;
            border: 1px solid #001E3C !important;
            font-weight: 500;
        }
        .stButton>button:hover {
            background-color: #004B8D !important;
            color: white !important;
        }

        .stDownloadButton>button {
            border-radius: 6px;
        }

        .stRadio>div, .stSelectbox>div {
            padding-top: 0.4rem;
            padding-bottom: 0.4rem;
        }

        .stTextArea textarea {
            border-radius: 6px !important;
        }

        .top-bar {
            background: linear-gradient(90deg, #002B5C, #004B8D);
            color: white;
            padding: 1rem 1.5rem;
            border-left: 6px solid #E87722;
            border-radius: 0 0 8px 8px;
            margin-bottom: 1.5rem;
        }
        .top-bar-title {
            font-size: 1.2rem;
            font-weight: 600;
            letter-spacing: 0.04em;
        }
        .top-bar-subtitle {
            font-size: 0.9rem;
            opacity: 0.95;
        }

        .footer {
            margin-top: 2rem;
            padding-top: 0.8rem;
            border-top: 1px solid #dde1e7;
            font-size: 0.8rem;
            color: #777;
            text-align: center;
        }
        </style>
        """,
        unsafe_allow_html=True,
    )


# ---------- Streamlit UI ----------

def main():
    st.set_page_config(
        page_title="AASU Research Proposal Evaluator",
        layout="wide",
        page_icon="📄",
    )

    inject_css()

    col_left, col_center, col_right = st.columns([1, 3, 1])

    with col_center:
        # Right-aligned logo
        st.markdown(
            """
            <div style="display: flex; justify-content: flex-end; margin-bottom: 0.5rem;">
                <img src="aasu_logo.jpg" width="130">
            </div>
            """,
            unsafe_allow_html=True,
        )

        # Language toggle
        language = st.radio(
            "Interface language / لغة الواجهة",
            options=["English", "Arabic"],
            index=0,
            horizontal=True,
        )

        # Main title block
        st.markdown(
            """
            <div style="text-align: center; margin-top: 0.3rem;">
                <h1 style="margin-bottom: 0.2rem; color: #002B5C;">
                    Abdullah Al-Salem University
                </h1>
                <h3 style="margin-top: 0; margin-bottom: 0.5rem; color: #E87722; font-weight: 500;">
                    AI Research Proposal Evaluation Tool
                </h3>
            </div>
            """,
            unsafe_allow_html=True,
        )

        if language == "Arabic":
            subtitle_text = (
                "أداة مساندة تعتمد على الذكاء الاصطناعي لدعم تقييم مقترحات البحث "
                "وفق الأولويات الوطنية والمؤسسية."
            )
        else:
            subtitle_text = (
                "AI-assisted internal tool to support evaluation of research proposals "
                "against national and institutional priorities."
            )

        st.markdown(
            f"""
            <div style="text-align: center;">
                <p style="color: #555; font-size: 0.95rem;">
                    {subtitle_text}
                </p>
            </div>
            """,
            unsafe_allow_html=True,
        )

        # Top ribbon "navbar"
        if language == "Arabic":
            top_title = "مكتب البحث والابتكار"
            top_sub = "فحص مبدئي آلي للمقترحات من حيث المواءمة الاستراتيجية، والأخلاقيات، وبناء القدرات."
        else:
            top_title = "Office of Research & Innovation"
            top_sub = "Automated screening of proposals for strategic alignment, ethics, and capacity building."

        st.markdown(
            f"""
            <div class="top-bar">
                <div class="top-bar-title">{top_title}</div>
                <div class="top-bar-subtitle">
                    {top_sub}
                </div>
            </div>
            """,
            unsafe_allow_html=True,
        )

        # Main content block with tabs
        st.markdown('<div class="main-block">', unsafe_allow_html=True)

        tab_labels = ["Evaluation", "About & Instructions"] if language == "English" else ["التقييم", "حول الأداة وتعليمات الاستخدام"]
        tab_eval, tab_about = st.tabs(tab_labels)

        with tab_eval:
            # Model settings, upload, choices, submit in a form
            if language == "Arabic":
                st.markdown("### إعدادات نموذج الذكاء الاصطناعي")
                model_label = "اختر نموذج اللغة (المحرّك):"
                model_help = "اختر النموذج بناءً على مستوى التفاصيل والسرعة المطلوبة."
                model_caption = (
                    "استخدم gpt-4o أو gpt-4.1 للتقارير المفصلة عالية الجودة، "
                    "واستخدم الإصدارات mini للفحص السريع."
                )
            else:
                st.markdown("### Model settings")
                model_label = "Select language model (engine):"
                model_help = "Choose a model according to the required level of detail and turnaround time."
                model_caption = (
                    "Use gpt-4o / gpt-4.1 for higher-quality narrative evaluations; "
                    "use the ‘mini’ variants for quick screening."
                )

            with st.form("evaluation_form"):
                model_name = st.selectbox(
                    model_label,
                    options=[
                        "gpt-4.1",
                        "gpt-4o",
                        "gpt-4o-mini",
                        "gpt-4.1-mini",
                    ],
                    index=2,
                    help=model_help,
                )
                st.caption(model_caption)

                st.markdown("---")

                # Upload field
                if language == "Arabic":
                    st.markdown("### ملف المقترح البحثي")
                    upload_label = "قم برفع المقترح البحثي (PDF أو DOCX أو TXT أو Markdown):"
                else:
                    st.markdown("### Proposal document")
                    upload_label = "Upload the research proposal (PDF, DOCX, TXT, or Markdown):"

                uploaded_file = st.file_uploader(
                    upload_label,
                    type=["pdf", "docx", "txt", "md"],
                )

                st.markdown("---")

                # Evaluation focus
                if language == "Arabic":
                    st.markdown("### مجال التقييم")
                    focus_label = "اختر محور التقييم الأساسي:"
                    focus_options = {
                        "kuwait": "مدى المواءمة مع رؤية الكويت 2035",
                        "aasu": "مدى المواءمة مع رؤية ورسالة الجامعة",
                        "ethics": "الاعتبارات والتبعات الأخلاقية",
                        "capacity": "بناء القدرات في جامعة عبدالله السالم",
                        "all": "جميع ما سبق (تقييم شامل)",
                    }
                    detail_label = "مستوى التفصيل المطلوب:"
                    detail_short = "ملخص قصير"
                    detail_detailed = "تقرير تفصيلي"
                    extra_label = "تعليمات إضافية للمقيّم (اختياري):"
                    extra_placeholder = (
                        "مثال: التركيز على أخلاقيات استخدام البيانات؛ التعليق على تدريب الطلاب؛ "
                        "اعتبار لجنة الكلية للجودة البحثية هي الجمهور المستهدف..."
                    )
                    submit_text = "إنشاء تقرير التقييم"
                else:
                    st.markdown("### Evaluation focus")
                    focus_label = "Select the primary focus of this evaluation:"
                    focus_options = {
                        "kuwait": "Alignment with Kuwait 2035 Vision",
                        "aasu": "Alignment with AASU vision & mission",
                        "ethics": "Ethical considerations & implications",
                        "capacity": "Capacity building at AASU",
                        "all": "All of the above",
                    }
                    detail_label = "Required level of detail:"
                    detail_short = "Short summary"
                    detail_detailed = "Detailed report"
                    extra_label = "Additional guidance for the reviewer (optional):"
                    extra_placeholder = (
                        "For example: focus on AI ethics; comment on student training components; "
                        "assume the audience is a College Research Committee, etc."
                    )
                    submit_text = "Generate evaluation"

                aspect_choice_key = st.radio(
                    focus_label,
                    options=list(focus_options.keys()),
                    format_func=lambda k: focus_options[k],
                    index=4,
                )

                detail_choice = st.radio(
                    detail_label,
                    options=["short", "detailed"],
                    format_func=lambda v: detail_short if v == "short" else detail_detailed,
                    index=0,
                )

                extra_instructions = st.text_area(
                    extra_label,
                    placeholder=extra_placeholder,
                    height=130,
                )

                st.markdown("---")

                submit_button = st.form_submit_button(
                    submit_text,
                    use_container_width=True,
                )

            # Handle submission
            if submit_button:
                if uploaded_file is None:
                    if language == "Arabic":
                        st.error("يرجى رفع ملف المقترح البحثي قبل طلب التقييم.")
                    else:
                        st.error("Please upload a proposal file before requesting an evaluation.")
                    st.markdown("</div>", unsafe_allow_html=True)
                    return

                with st.spinner("..."):

                    proposal_text = read_uploaded_file(uploaded_file)

                    if not proposal_text.strip():
                        if language == "Arabic":
                            st.error(
                                "تعذر على النظام استخراج النص من الملف المرفوع. "
                                "يرجى تجربة صيغة أخرى (مثل DOCX أو ملف PDF نصي)."
                            )
                        else:
                            st.error(
                                "The system could not extract text from the uploaded file. "
                                "Please try a different format (for example, a DOCX or text-based PDF)."
                            )
                        st.markdown("</div>", unsafe_allow_html=True)
                        return

                    aspect_instructions = build_aspect_instructions(aspect_choice_key, language)
                    detail_instruction = build_detail_instruction(detail_choice, language)
                    language_instruction = build_language_instruction(language)

                    if language == "Arabic":
                        system_prompt = (
                            "You are an experienced academic reviewer at Abdullah Al-Salem University (AASU). "
                            "You evaluate research proposals for alignment with institutional strategies, "
                            "Kuwait national priorities, ethics, and capacity building.\n\n"
                            "Respond in formal Arabic suitable for academic committees. "
                            "Always be constructive, specific, and professional. When information is missing, "
                            "explicitly state what is missing rather than inventing details."
                        )
                    else:
                        system_prompt = (
                            "You are an experienced academic reviewer at Abdullah Al-Salem University (AASU). "
                            "You evaluate research proposals for alignment with institutional strategies, "
                            "Kuwait national priorities, ethics, and capacity building.\n\n"
                            "Respond in formal British English suitable for academic committees. "
                            "Always be constructive, specific, and professional. When information is missing, "
                            "explicitly state what is missing rather than inventing details. Be cretical and add score out of 100% for each section with justification"
                        )

                    template = ChatPromptTemplate.from_messages(
                        [
                            ("system", system_prompt),
                            (
                                "user",
                                (
                                    "You are given a research proposal. Your task:\n"
                                    "{aspect_instructions}\n\n"
                                    "{detail_instruction}\n\n"
                                    "{language_instruction}\n\n"
                                    "If extra instructions are provided, follow them as well:\n"
                                    "Extra instructions from the user: {extra_instructions}\n\n"
                                    "Here is the proposal text:\n"
                                    "-------------------------\n"
                                    "{proposal_text}\n"
                                    "-------------------------\n"
                                ),
                            ),
                        ]
                    )

                    llm = get_llm(model_name)
                    chain = template | llm

                    response = chain.invoke(
                        {
                            "aspect_instructions": aspect_instructions,
                            "detail_instruction": detail_instruction,
                            "language_instruction": language_instruction,
                            "extra_instructions": extra_instructions or "None provided.",
                            "proposal_text": proposal_text,
                        }
                    )

                    evaluation_text = response.content

                # Show evaluation
                if language == "Arabic":
                    st.markdown("### ناتج التقييم")
                    st.markdown(
                        "التقرير الآتي يعد مسودة مبدئية لمساعدة المقيّمين، "
                        "وينبغي قراءته بالتوازي مع المقترح الكامل."
                    )
                else:
                    st.markdown("### Evaluation output")
                    st.markdown(
                        "The following draft evaluation is intended as an internal aid for reviewers "
                        "and should be read alongside the full proposal."
                    )

                st.markdown("---")
                st.markdown(evaluation_text)

                # Download report
                st.markdown("---")
                if language == "Arabic":
                    st.markdown("### تحميل تقرير التقييم")
                    dl_label = "تحميل تقرير Word (.docx)"
                    detail_label_text = detail_short if detail_choice == "short" else detail_detailed
                    aspect_label_text = focus_options[aspect_choice_key]
                else:
                    st.markdown("### Download evaluation report")
                    dl_label = "Download Word report (.docx)"
                    detail_label_text = "Short summary" if detail_choice == "short" else "Detailed report"
                    aspect_label_text = focus_options[aspect_choice_key]

                report_buffer = create_report_docx(
                    proposal_filename=uploaded_file.name,
                    model_name=model_name,
                    aspect_label=aspect_label_text,
                    detail_label=detail_label_text,
                    extra_instructions=extra_instructions,
                    evaluation_text=evaluation_text,
                    language=language,
                )

                timestamp = datetime.now().strftime("%Y%m%d_%H%M")
                default_filename = f"AASU_Proposal_Evaluation_{timestamp}.docx"

                st.download_button(
                    label=dl_label,
                    data=report_buffer,
                    file_name=default_filename,
                    mime=(
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    ),
                    use_container_width=True,
                )

        with tab_about:
            # About & instructions tab
            if language == "Arabic":
                st.markdown("### حول الأداة")
                st.write(
                    "تهدف هذه الأداة إلى دعم لجان البحث في جامعة عبدالله السالم من خلال "
                    "توليد مسودات تقييم آلية للمقترحات البحثية، مع التركيز على المواءمة "
                    "مع رؤية الكويت 2035، ورؤية ورسالة الجامعة، والجوانب الأخلاقية، "
                    "وبناء القدرات داخل الجامعة."
                )
                st.markdown("### طريقة الاستخدام")
                st.markdown(
                    """
                    1. اختر نموذج الذكاء الاصطناعي (المحرّك) المناسب.
                    2. ارفع ملف المقترح البحثي بصيغة PDF أو DOCX أو TXT أو Markdown.
                    3. حدد محور التقييم الأساسي أو اختر “جميع ما سبق” لتقييم شامل.
                    4. اختر مستوى التفصيل (ملخص قصير أو تقرير تفصيلي).
                    5. (اختياري) أضف أي تعليمات إضافية مثل التركيز على جانب معين.
                    6. اضغط زر “إنشاء تقرير التقييم” للحصول على المسودة.
                    7. يمكنك تنزيل التقييم كملف Word لاستخدامه في لجان الكلية أو الجامعة.
                    """
                )
                st.markdown("### ملاحظة مهمة")
                st.write(
                    "الناتج الذي تولده الأداة يعد مسودة مساعدة فقط، ولا يغني عن الحكم الأكاديمي "
                    "المباشر من أعضاء الهيئة التدريسية واللجان المختصة."
                )
            else:
                st.markdown("### About this tool")
                st.write(
                    "This application supports research committees at Abdullah Al-Salem University "
                    "by generating AI-assisted draft evaluations of research proposals. "
                    "It focuses on alignment with Kuwait Vision 2035, AASU’s vision and mission, "
                    "ethical considerations, and capacity building within the university."
                )
                st.markdown("### How to use")
                st.markdown(
                    """
                    1. Select the preferred language model (engine).
                    2. Upload the research proposal as a PDF, DOCX, TXT, or Markdown file.
                    3. Choose the primary evaluation focus, or select “All of the above” for a holistic review.
                    4. Choose the required level of detail (short summary or detailed report).
                    5. Optionally provide additional instructions, such as emphasising ethics or student training.
                    6. Click “Generate evaluation” to produce the draft.
                    7. Download the evaluation as a Word (.docx) report for use in School / College / University committees.
                    """
                )
                st.markdown("### Important note")
                st.write(
                    "The generated output is intended as an internal aid and does not replace "
                    "direct academic judgement by faculty members and formal review committees."
                )

        st.markdown("</div>", unsafe_allow_html=True)  # close main-block

        # Footer
        if language == "Arabic":
            footer_text = (
                "© Abdullah Al-Salem University – مكتب البحث والابتكار\n"
                "يتم إنشاء هذا التقييم باستخدام نماذج ذكاء اصطناعي، "
                "ويجب قراءته كأداة مساندة وليس كحكم نهائي."
            )
        else:
            footer_text = (
                "© Abdullah Al-Salem University – Office of Research & Innovation\n"
                "This evaluation is generated using AI models and should be treated as decision support, "
                "not as a final judgement."
            )

        st.markdown(
            f"""
            <div class="footer">
                {footer_text}
            </div>
            """,
            unsafe_allow_html=True,
        )


if __name__ == "__main__":
    main()
